import os
import re
import json
from pathlib import Path
from typing import List, Dict, Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS


DATA_DIR = "data"
VECTOR_DIR = "vectorstore/faiss_fit_hcmute"


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[-=]{5,}", " ", text)
    text = re.sub(r"\n+", "\n", text)
    return text.strip()


def guess_category(filename: str) -> str:
    name = filename.lower()

    if "tuyen" in name or "admission" in name or "phuongthuc" in name or "deantuyensinh" in name:
        return "admission"

    if "nganh" in name or "chuongtrinh" in name or "daotao" in name or "cntt" in name or "kmt" in name:
        return "curriculum"

    if "giangvien" in name or "doi ngu" in name or "doingugiangvien" in name:
        return "lecturers"

    if "nghiencuu" in name or "research" in name:
        return "research"

    if "thongbao" in name or "event" in name or "sukien" in name:
        return "announcements"

    if "gioithieu" in name or "tamnhin" in name or "sumang" in name or "chucnang" in name or "tochuc" in name:
        return "faculty_info"

    return "general"


def read_txt(path: Path) -> str:
    encodings = ["utf-8", "utf-8-sig", "cp1258", "latin-1"]

    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except Exception:
            continue

    return ""


def read_docx(path: Path) -> str:
    try:
        doc = DocxDocument(str(path))
        parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
    except Exception:
        return ""


def read_pdf(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)

        return "\n".join(parts)
    except Exception:
        return ""


def normalize_json_item(item: Any, filename: str) -> List[Document]:
    docs = []

    if isinstance(item, dict):
        title = str(
            item.get("title")
            or item.get("TIÊU ĐỀ")
            or item.get("tieu_de")
            or item.get("name")
            or filename
        )

        url = str(
            item.get("url")
            or item.get("URL")
            or item.get("link")
            or ""
        )

        date = str(
            item.get("date")
            or item.get("NGÀY")
            or item.get("created_date")
            or ""
        )

        content = str(
            item.get("content")
            or item.get("text")
            or item.get("main_content")
            or item.get("body")
            or item.get("NỘI DUNG")
            or ""
        )

        if not content.strip():
            content_parts = []
            for key, value in item.items():
                if isinstance(value, str):
                    content_parts.append(f"{key}: {value}")
            content = "\n".join(content_parts)

        content = clean_text(content)

        if len(content) >= 30:
            docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "title": title,
                        "source": filename,
                        "url": url,
                        "date": date,
                        "category": guess_category(filename),
                        "source_type": "json",
                    },
                )
            )

    elif isinstance(item, str):
        content = clean_text(item)
        if len(content) >= 30:
            docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "title": filename,
                        "source": filename,
                        "url": "",
                        "date": "",
                        "category": guess_category(filename),
                        "source_type": "json",
                    },
                )
            )

    return docs


def read_json(path: Path) -> List[Document]:
    try:
        text = read_txt(path)
        data = json.loads(text)
        filename = path.name
        docs = []

        if isinstance(data, list):
            for item in data:
                docs.extend(normalize_json_item(item, filename))

        elif isinstance(data, dict):
            if "data" in data and isinstance(data["data"], list):
                for item in data["data"]:
                    docs.extend(normalize_json_item(item, filename))
            elif "articles" in data and isinstance(data["articles"], list):
                for item in data["articles"]:
                    docs.extend(normalize_json_item(item, filename))
            else:
                docs.extend(normalize_json_item(data, filename))

        return docs
    except Exception:
        return []


def load_documents(data_dir: str) -> List[Document]:
    documents = []
    paths = list(Path(data_dir).rglob("*"))

    for path in paths:
        if not path.is_file():
            continue

        suffix = path.suffix.lower()
        filename = path.name

        if suffix == ".json":
            documents.extend(read_json(path))
            continue

        if suffix == ".txt":
            content = read_txt(path)
            source_type = "txt"
        elif suffix == ".docx":
            content = read_docx(path)
            source_type = "docx"
        elif suffix == ".pdf":
            content = read_pdf(path)
            source_type = "pdf"
        else:
            continue

        content = clean_text(content)

        if len(content) < 30:
            continue

        documents.append(
            Document(
                page_content=content,
                metadata={
                    "title": path.stem,
                    "source": filename,
                    "url": "",
                    "date": "",
                    "category": guess_category(filename),
                    "source_type": source_type,
                },
            )
        )

    return documents


def split_documents(documents: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=150,
        separators=["\n\n", "\n", ".", "!", "?", ";", ",", " ", ""],
    )

    chunks = splitter.split_documents(documents)

    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = i

    return chunks


def main():
    documents = load_documents(DATA_DIR)

    print(f"Số tài liệu gốc: {len(documents)}")

    chunks = split_documents(documents)

    print(f"Số chunk sau khi chia: {len(chunks)}")

    embeddings = HuggingFaceEmbeddings(
        model_name="intfloat/multilingual-e5-base",
        encode_kwargs={"normalize_embeddings": True},
    )

    vectorstore = FAISS.from_documents(chunks, embeddings)

    os.makedirs(VECTOR_DIR, exist_ok=True)

    vectorstore.save_local(VECTOR_DIR)

    print(f"Đã lưu vector database vào: {VECTOR_DIR}")


if __name__ == "__main__":
    main()